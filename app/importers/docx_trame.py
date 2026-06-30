"""Import d'une trame depuis un .docx (US1.3).

Deux étages, séparés pour la testabilité :
- ``extract_blocks(doc)`` : .docx -> liste de ``Block`` (niveau, texte), en
  s'appuyant sur les niveaux de liste Word (``w:ilvl``) et les styles Heading.
- ``build_trame(blocks)`` : ``Block[]`` -> ``ParsedTrame`` (thèmes/questions
  aplatis sur 2 niveaux). Logique pure, testable sans fichier.

Convention reconnue (format type assessment OCTO) :
  niveau 0 numéroté = thème · puce niveau 1 = question ·
  sous-puce niveau 2 = sous-question (remontée) ou options (-> question 'choix').

Outil d'inspection (pour caler le parser sur un vrai document) :
    python -m app.importers.docx_trame inspect <fichier.docx>
    python -m app.importers.docx_trame <fichier.docx>      # aperçu du résultat
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field

from docx import Document
from docx.oxml.ns import qn

ANCHOR = "interview"
_NUM_RE = re.compile(r"^\s*\d+[.)]\s+")
_BULLET_RE = re.compile(r"^[•▪●◦‣o]\s+")
# Section d'introduction d'une trame (briefing avant les questions).
_INTRO_RE = re.compile(r"objectifs?\s+et\s+principes", re.IGNORECASE)


@dataclass
class Block:
    level: int
    text: str
    numbered: bool = False
    style: str = ""


@dataclass
class ParsedQuestion:
    label: str
    qtype: str = "open"
    config: dict = field(default_factory=dict)
    help: str = ""


@dataclass
class ParsedTheme:
    title: str
    questions: list[ParsedQuestion] = field(default_factory=list)


@dataclass
class ParsedTrame:
    name: str
    themes: list[ParsedTheme] = field(default_factory=list)
    # Texte d'introduction « Objectifs et principes » (vide si absent).
    intro: str = ""


def _strip_marker(text: str) -> str:
    """Retire un éventuel préfixe de liste résiduel (numéro / puce)."""
    text = _NUM_RE.sub("", text)
    text = _BULLET_RE.sub("", text)
    return text.strip()


# --------------------------------------------------------------------------- #
# Étage 1 : .docx -> Block[]
# --------------------------------------------------------------------------- #
def _ilvl(paragraph) -> int | None:
    """Niveau de liste Word d'un paragraphe (None s'il n'est pas en liste)."""
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    ilvl = numPr.find(qn("w:ilvl"))
    return int(ilvl.get(qn("w:val"))) if ilvl is not None else 0


def _level_of(paragraph, text: str) -> tuple[int, bool]:
    ilvl = _ilvl(paragraph)
    if ilvl is not None:
        return ilvl, bool(_NUM_RE.match(text))
    style = paragraph.style.name or ""
    if style.startswith("Heading"):
        try:
            return max(0, int(style.split()[-1]) - 1), False
        except ValueError:
            return 0, False
    if _NUM_RE.match(text):
        return 0, True
    # Paragraphe hors-liste : un intertitre (contient « : » sans « ? ») ouvre
    # un thème (ex. "Processus budgétaire :", "Portefeuille de projets : …").
    if ":" in text and "?" not in text:
        return 0, False
    return 1, False


def extract_blocks(doc, anchor: str = ANCHOR) -> list[Block]:
    paras = [p for p in doc.paragraphs if p.text.strip()]

    # Point de départ : l'ancre "Interview" (saute le briefing), sinon le 1er
    # item numéroté rencontré.
    start = 0
    for i, p in enumerate(paras):
        if p.text.strip().lower() == anchor:
            start = i + 1
            break
    else:
        for i, p in enumerate(paras):
            if _NUM_RE.match(p.text.strip()):
                start = i
                break

    blocks: list[Block] = []
    for p in paras[start:]:
        text = p.text.strip()
        level, numbered = _level_of(p, text)
        blocks.append(
            Block(level=level, text=text, numbered=numbered, style=p.style.name or "")
        )
    return blocks


def extract_intro(doc, anchor: str = ANCHOR) -> str:
    """Texte d'introduction « Objectifs et principes ».

    Repère la section dont le titre contient « objectifs et principes » et
    collecte les paragraphes qui suivent, jusqu'à l'ancre « Interview » ou au
    premier item numéroté (= début des questions). Chaîne vide si absent.
    """
    paras = [p for p in doc.paragraphs if p.text.strip()]
    start = None
    for i, p in enumerate(paras):
        if _INTRO_RE.search(p.text):
            start = i
            break
    if start is None:
        return ""

    out: list[str] = []
    for offset, p in enumerate(paras[start:]):
        text = p.text.strip()
        if offset > 0:  # on garde toujours la ligne de titre, puis on borne
            if text.lower() == anchor.lower():
                break
            if _NUM_RE.match(text):
                break
        out.append(text)
    return "\n".join(out).strip()


# --------------------------------------------------------------------------- #
# Étage 2 : Block[] -> ParsedTrame  (logique pure)
# --------------------------------------------------------------------------- #
def build_trame(blocks: list[Block], name: str = "Trame importée") -> ParsedTrame:
    """Block[] -> ParsedTrame.

    Règles :
    - niveau 0 = thème ;
    - toute ligne se terminant par « ? » = question ;
    - les autres lignes sont du **texte d'aide** : soit une amorce qui introduit
      les questions suivantes (typiquement terminée par « : »), bufferisée et
      rattachée à la prochaine question ; soit une précision/exemple qui suit
      une question, rattachée à la question courante. (On ne fabrique plus de
      fausses « options » à partir de ces lignes.)
    """
    themes: list[ParsedTheme] = []
    current: ParsedTheme | None = None
    current_q: ParsedQuestion | None = None
    pending_help: list[str] = []

    def attach(q: ParsedQuestion) -> None:
        nonlocal pending_help
        if pending_help:
            extra = "\n".join(pending_help)
            q.help = f"{q.help}\n{extra}".strip() if q.help else extra
            pending_help = []

    for b in blocks:
        text = _strip_marker(b.text)
        if not text or not re.search(r"\w", text):  # lignes ponctuation (« : », …)
            continue

        if b.level <= 0:
            current = ParsedTheme(title=text)
            themes.append(current)
            current_q = None
            pending_help = []
        elif text.endswith("?"):
            if current is None:
                current = ParsedTheme(title="Général")
                themes.append(current)
            current_q = ParsedQuestion(label=text)
            attach(current_q)  # l'amorce qui précède devient l'aide
            current.questions.append(current_q)
        else:
            # Texte d'aide. Amorce (« : ») ou bloc d'amorce déjà commencé, ou pas
            # encore de question -> on bufferise pour la prochaine question.
            # Sinon -> précision rattachée à la question courante.
            if pending_help or text.endswith(":") or current_q is None:
                pending_help.append(text)
            else:
                current_q.help = (
                    f"{current_q.help}\n{text}".strip() if current_q.help else text
                )

    themes = [t for t in themes if t.questions]
    return ParsedTrame(name=name, themes=themes)


# --------------------------------------------------------------------------- #
# Façade
# --------------------------------------------------------------------------- #
def parse_docx(source, name: str = "Trame importée") -> ParsedTrame:
    """`source` : chemin ou objet fichier (.docx)."""
    doc = Document(source)
    trame = build_trame(extract_blocks(doc), name=name)
    trame.intro = extract_intro(doc)
    return trame


def parse_docx_bytes(content: bytes, name: str = "Trame importée") -> ParsedTrame:
    return parse_docx(io.BytesIO(content), name=name)


# --------------------------------------------------------------------------- #
# Inspection / CLI (pour caler le parser sur un vrai document)
# --------------------------------------------------------------------------- #
def inspect(path: str) -> None:
    doc = Document(path)
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        num = bool(_NUM_RE.match(t))
        print(f"ilvl={_ilvl(p)!s:>4}  style={p.style.name!r:<16} num={num!s:<5} | {t[:90]}")


if __name__ == "__main__":
    import sys

    if len(sys.argv) >= 3 and sys.argv[1] == "inspect":
        inspect(sys.argv[2])
    elif len(sys.argv) >= 2:
        trame = parse_docx(sys.argv[1])
        print(f"Trame : {trame.name}  ({len(trame.themes)} thèmes)")
        for th in trame.themes:
            print(f"\n# {th.title}")
            for q in th.questions:
                extra = f"   [{q.qtype}: {q.config}]" if q.qtype != "open" else ""
                print(f"  - {q.label}{extra}")
                if q.help:
                    for line in q.help.splitlines():
                        print(f"      · aide: {line}")
    else:
        print(__doc__)
